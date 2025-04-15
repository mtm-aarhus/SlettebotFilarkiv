"""This module contains the main process of the robot."""

from OpenOrchestrator.orchestrator_connection.connection import OrchestratorConnection
from OpenOrchestrator.database.queues import QueueElement
import re
import os
import pandas as pd
from office365.runtime.auth.user_credential import UserCredential
from office365.sharepoint.client_context import ClientContext
from docx.shared import Pt
from docx.shared import Inches
from docx import Document
import json
import zipfile
import shutil
from datetime import date
import datetime
import xml.etree.ElementTree as ET
import AfslutSag
import GetKmdAcessToken
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# pylint: disable-next=unused-argument
def process(orchestrator_connection: OrchestratorConnection, queue_element: QueueElement | None = None) -> None:
    print('I gang')
    """This module contains the main process of the robot."""
    def sharepoint_client(username, password, sharepoint_site_url):
        ctx = ClientContext(sharepoint_site_url).with_credentials(UserCredential(username, password))
        web = ctx.web
        ctx.load(web)
        ctx.execute_query()
        return ctx

    def upload_to_sharepoint(client: ClientContext, folder_name: str, file_path: str, folder_url: str):
        """
        Uploads a file to a specific folder in a SharePoint document library.

        :param client: Authenticated SharePoint client context
        :param folder_name: Name of the target folder within the document library
        :param file_path: Local file path to upload
        :param folder_url: SharePoint folder URL where the file should be uploaded
        """
        try:
            # Extract file name safely
            file_name = os.path.basename(file_path)

            # Define the SharePoint document library structure
            document_library = f"{folder_url.split('/', 1)[-1]}/Delte Dokumenter/Aktindsigter"
            folder_path = f"{document_library}/{folder_name}"

            # Read file into memory (Prevents closed file issue)
            with open(file_path, "rb") as file:
                file_content = file.read()  

            # Get SharePoint folder reference
            target_folder = client.web.get_folder_by_server_relative_url(folder_url)

            # Upload file using byte content
            target_folder.upload_file(file_name, file_content)
            
            # Execute request
            client.execute_query()
            orchestrator_connection.log_info(f"✅ Successfully uploaded: {file_name} to {folder_path}")

        except Exception as e:
            orchestrator_connection.log_info(f"❌ Error uploading file: {str(e)}")

    def download_file_from_sharepoint(client, sharepoint_file_url):
        '''
        Function for downloading file from sharepoint
        '''
        file_name = sharepoint_file_url.split("/")[-1]
        download_path = os.path.join(os.getcwd(), file_name)
        with open(download_path, "wb") as local_file:
            client.web.get_file_by_server_relative_path(sharepoint_file_url).download(local_file).execute_query()
        return download_path

    def check_excel_file(file_path):
        '''
        Goes through the document list and saves the data in a dictionary.
        '''
        df = pd.read_excel(file_path)
        documents = []
        if 'Gives der aktindsigt i dokumentet? (Ja/Nej/Delvis)' in df.columns and 'Begrundelse hvis nej eller delvis' in df.columns:
            for _, row in df.iterrows():
                documents.append({
                    'title': row['Dokumenttitel'],
                    'decision': row['Gives der aktindsigt i dokumentet? (Ja/Nej/Delvis)'],
                    'reason': row['Begrundelse hvis nej eller delvis'],
                    'Akt ID': row['Akt ID'],
                    'Dok ID': row['Dok ID']
                })
        return documents

    def traverse_and_check_folders(client, folder_url, results, orchestrator_connection):
        '''
        Goes through the different folders to find the excel file (ie. the document list)
        '''
        pattern = re.compile(r"([A-Za-z]\d{4}-\d{1,10}|[A-Za-z]{3}-\d{4}-\d{6})")
        folder = client.web.get_folder_by_server_relative_url(folder_url)
        client.load(folder)
        client.execute_query()

        subfolders = folder.folders
        client.load(subfolders)
        client.execute_query()

        for subfolder in subfolders:
            subfolder_name = subfolder.properties["Name"]
            subfolder_url = f"{folder_url}/{subfolder_name}"
            if re.search(pattern, subfolder_name):
                files = subfolder.files
                client.load(files)
                client.execute_query()

                for file in files:
                    if file.properties["Name"].endswith(".xlsx"):
                        file_url = f"{subfolder_url}/{file.properties['Name']}"
                        local_file_path = download_file_from_sharepoint(client, file_url)
                        document_results = check_excel_file(local_file_path)
                        results[subfolder_name] = document_results  # Ensuring it is a list
                        os.remove(local_file_path)
                        break

            traverse_and_check_folders(client, subfolder_url, results, orchestrator_connection)

    def replace_placeholders_in_xml(docx_path: str, replacements: dict):
        temp_dir = "temp_xml_unzip"
        unzip_path = os.path.join(temp_dir, "unzipped")
        os.makedirs(unzip_path, exist_ok=True)

        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(unzip_path)

        word_folder = os.path.join(unzip_path, "word")
        targets = [
            f for f in os.listdir(word_folder)
            if f.startswith(("document", "header", "footer")) and f.endswith(".xml")
        ]

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        for filename in targets:
            xml_path = os.path.join(word_folder, filename)
            tree = ET.parse(xml_path)
            root = tree.getroot()

            # Gennemgå alle afsnit <w:p>
            for para in root.findall('.//w:p', ns):
                runs = para.findall('.//w:r', ns)
                full_text = ""
                text_nodes = []

                for run in runs:
                    for t in run.findall('.//w:t', ns):
                        text_nodes.append(t)
                        full_text += t.text if t.text else ""

                replaced_text = full_text
                for ph, val in replacements.items():
                    replaced_text = replaced_text.replace(ph, val)

                if replaced_text != full_text:
                    # Slet eksisterende tekstindhold
                    for t in text_nodes:
                        t.text = ""

                    # Fordel ny tekst i samme struktur
                    remaining = replaced_text
                    for t in text_nodes:
                        if not remaining:
                            break
                        t.text = remaining[:len(remaining)]
                        remaining = ""

            tree.write(xml_path, encoding='utf-8', xml_declaration=True)

        # Zip tilbage
        new_docx_path = docx_path.replace(".docx", "_updated.docx")
        with zipfile.ZipFile(new_docx_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for foldername, _, filenames in os.walk(unzip_path):
                for filename in filenames:
                    filepath = os.path.join(foldername, filename)
                    arcname = os.path.relpath(filepath, unzip_path)
                    zipf.write(filepath, arcname)

        shutil.rmtree(temp_dir)
        return new_docx_path

    def insert_list_at_placeholder(doc, placeholder, case_details, fontsize=9):
        full_access_cases = []
        limited_access_cases = []

        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.clear()
                insert_index = paragraph._element
                parent = paragraph._element.getparent()
                insert_position = parent.index(paragraph._element)

                for case_title, documents in case_details.items():
                    filtered_docs = [doc for doc in documents if doc['decision'] in ['Nej', 'Delvis']]
                    if not filtered_docs:
                        full_access_cases.append(f"• {case_title}")
                    else:
                        limited_access_cases.append(case_title)

                if full_access_cases:
                    p = doc.add_paragraph("Der gives fuld aktindsigt i følgende sager:")
                    p.runs[0].bold = True
                    parent.insert(insert_position + 1, p._element)
                    insert_position += 1

                    p = doc.add_paragraph("\n".join(full_access_cases))
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.runs[0].font.size = Pt(fontsize)
                    parent.insert(insert_position + 1, p._element)
                    insert_position += 1

                if limited_access_cases:
                    p = doc.add_paragraph("\nDer gives delvis eller ingen aktindsigt i følgende sager:")
                    p.runs[0].bold = True
                    p.paragraph_format.space_after = Pt(5)
                    parent.insert(insert_position + 1, p._element)
                    insert_position += 1

                    for case_title in limited_access_cases:
                        p = doc.add_paragraph(f"• {case_title}")
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.runs[0].font.size = Pt(fontsize)
                        parent.insert(insert_position + 1, p._element)
                        insert_position += 1

                        filtered_docs = [doc for doc in case_details[case_title] if doc['decision'] in ['Nej', 'Delvis']]

                        if len(filtered_docs) > 10:
                            p = doc.add_paragraph("Der er mange filer i denne sag. Se aktlisten for overblik over de enkelte filer.")
                            p.paragraph_format.left_indent = Inches(0.5)
                            p.runs[0].font.size = Pt(fontsize)
                            parent.insert(insert_position + 1, p._element)
                            insert_position += 1
                        else:
                            for document in filtered_docs:
                                reason_text = document['reason'] if len(str(document['reason'])) > 3 else "Ingen yderligere begrundelse"
                                akt_id_formatted = str(document["Akt ID"]).zfill(4)

                                p = doc.add_paragraph("• ")
                                p.paragraph_format.left_indent = Inches(0.5)
                                p.paragraph_format.space_after = Pt(0)

                                p.add_run(f"{akt_id_formatted}-{document['Dok ID']}, ").font.size = Pt(fontsize)

                                r = p.add_run("Aktindsigt:")
                                r.italic = True
                                r.font.size = Pt(fontsize)

                                p.add_run(f" {document['decision']}, ").font.size = Pt(fontsize)

                                r = p.add_run("Begrundelse:")
                                r.italic = True
                                r.font.size = Pt(fontsize)

                                p.add_run(f" {reason_text}").font.size = Pt(fontsize)

                                parent.insert(insert_position + 1, p._element)
                                insert_position += 1
                break

    def insert_table_at_placeholder(doc, placeholder, case_details, fontsize=9):
        print('Running insert table at placeholder')
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                # Find forælder og indsættelsesposition
                parent = paragraph._element.getparent()
                insert_position = parent.index(paragraph._element)

                # Byg data til tabellen
                table_data = []
                for case_title, documents in case_details.items():
                    decisions = [doc['decision'] for doc in documents]

                    if all(d == 'Ja' for d in decisions):
                        status = "Fuld aktindsigt"
                    elif all(d == 'Nej' for d in decisions):
                        status = "Ingen aktindsigt"
                    else:
                        status = "Delvis aktindsigt"

                    table_data.append((case_title, status))

                # Opret tabel
                print('Making table')
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'

                hdr_cells = table.rows[0].cells
                headers = ["Sagsnavn", "Fuld, delvis eller ingen aktindsigt"]

                for i, text in enumerate(headers):
                    p = hdr_cells[i].paragraphs[0]
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(fontsize)

                    # Tilføj grå baggrund
                    tc = hdr_cells[i]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), "D9D9D9")  # Lys grå
                    tcPr.append(shd)

                # Tilføj data
                for case_title, status in table_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = case_title
                    row_cells[1].text = status

                    for cell in row_cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(fontsize)
                # Indsæt tabel før vi fjerner placeholder-paragraf
                parent.insert(insert_position, table._element)
                print('Table inserted')

                # Fjern den gamle placeholder-paragraf
                parent.remove(paragraph._element)
                break

    def update_document_with_besvarelse(doc_path, case_details, DeskproTitel, AnsøgerNavn, AnsøgerEmail, Afdeling, AktindsigtsDato):
        doc = Document(doc_path)
        insert_table_at_placeholder(doc, "[Sagstabel]", case_details)
        temp_path = "Afgørelsesskriv.docx"
        doc.save(temp_path)

        replacements = {
            "[Deskprotitel]": DeskproTitel,
            "[Ansøgernavn]": AnsøgerNavn,
            "[Ansøgermail]": AnsøgerEmail,
            "[Afdeling]": Afdeling,
            "[Modtagelsesdato]": datetime.datetime.strptime(AktindsigtsDato, "%Y-%m-%dT%H:%M:%SZ").strftime("%d-%m-%Y")
        }

        updated_path = replace_placeholders_in_xml(temp_path, replacements)
        os.replace(updated_path, temp_path)


    queue_json = json.loads(queue_element.data)
    DeskproTitel = queue_json.get('Aktindsigtsovermappe')
    AnsøgerNavn = queue_json.get('AnsøgerNavn')
    AnsøgerEmail = queue_json.get('AnsøgerEmail')
    Afdeling = queue_json.get('Afdeling')
    DeskProID = queue_json.get('DeskProID')
    KMDNovaURL = orchestrator_connection.get_constant("KMDNovaURL").value
    AktindsigtsDato = queue_json.get("AktindsigtsDato")

    orchestrator_connection.log_info(f'processing {DeskproTitel}')

    orchestrator_connection.log_info('Getting credentials')
    RobotCredentials = orchestrator_connection.get_credential("RobotCredentials")
    username = RobotCredentials.username
    password = RobotCredentials.password
    sharepoint_site_url = orchestrator_connection.get_constant("AktbobSharePointURL").value
    parent_folder_url = sharepoint_site_url.split(".com")[-1] +'/Delte Dokumenter/'

    orchestrator_connection.log_info('Getting client')
    client = sharepoint_client(username, password, sharepoint_site_url)
    results = {}
    orchestrator_connection.log_info('Going through folder')
    traverse_and_check_folders(client, f'{parent_folder_url}Dokumentlister/{DeskproTitel}', results, orchestrator_connection)
    doc_path = r'Document.docx'
    orchestrator_connection.log_info('Updating document')
    update_document_with_besvarelse(doc_path, results, DeskproTitel= DeskproTitel, AnsøgerEmail= AnsøgerEmail, AnsøgerNavn= AnsøgerNavn, Afdeling= Afdeling, AktindsigtsDato = AktindsigtsDato)
    orchestrator_connection.log_info('Setting cases as finished in nova if novacase')
    KMD_access_token = GetKmdAcessToken.GetKMDToken(orchestrator_connection = orchestrator_connection)
    AfslutSag.invoke_AfslutSag(KMDNovaURL, KMD_access_token, DeskProID= DeskProID, orchestrator_connection= orchestrator_connection)
    orchestrator_connection.log_info('Document updating, uploading to sharepoint')
    upload_to_sharepoint(client, DeskproTitel, r'Afgørelsesskriv.docx', folder_url = f'{parent_folder_url}Aktindsigter/{DeskproTitel}')