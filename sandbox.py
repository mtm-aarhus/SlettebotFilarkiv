from docx import Document
import requests
from OpenOrchestrator.orchestrator_connection.connection import OrchestratorConnection
import os
from requests_ntlm import HttpNtlmAuth
import xml.etree.ElementTree as ET
from datetime import datetime

deskproid = "2076"


orchestrator_connection = OrchestratorConnection("AktindsigtAfgørelsesskriv", os.getenv('OpenOrchestratorSQL'),os.getenv('OpenOrchestratorKey'), None)
aktbob_credentials = orchestrator_connection.get_credential("AktbobAPIKey")
base_url = aktbob_credentials.username
url = f"{base_url}/Database/Tickets?deskproId={deskproid}"
headers = {
  'ApiKey': aktbob_credentials.password
}

response = requests.request("GET", url, headers=headers)
data = response.json()
# Extracting caseNumber values
case_numbers = [
    case["caseNumber"] for case in data[0]["cases"] 
    if case["sharepointFolderName"] is not None
]

if case_numbers:
    case_details = []  # List to hold each case's details
    go_credentials = orchestrator_connection.get_credential("GOAktApiUser")
    API_url = orchestrator_connection.get_constant("GOApiURL").value
    session = requests.Session()
    session.auth = HttpNtlmAuth(go_credentials.username, go_credentials.password)
    session.post(API_url, timeout=500)
    for case in case_numbers:
        response = session.get(f'{API_url}/_goapi/Cases/Metadata/{case}')
        data = response.json()
        metadata_xml = data["Metadata"]
        # Parse the XML and fetch the ows_Title attribute
        root = ET.fromstring(metadata_xml)
        case_title = root.get("ows_Title")
        modtaget_date = datetime.strptime(root.get("ows_Modtaget"), "%Y-%m-%d %H:%M:%S").strftime("%d-%m-%Y")
        aktindsigt_decision = "Your Aktindsigt Decision Here"  # Customize this as needed

        # Add the details to the list
        case_details.append([case, case_title, modtaget_date, aktindsigt_decision])
        
        
        
            
print(case_numbers)

# Load the document
doc = Document('Document.docx')

# Define the variables for each unique placeholder
afdeling = "Digitalisering"
ansoegernavn = "John Doe"
ansoegermail = "john.doe@example.com"
dato = "11. november 2024"
deskprotitel = "Ejendomssag"
besvarelse = "Din anmodning er blevet godkendt."
afdelingsmail = "digitalisering@aarhus.dk"
afdelingstelefon = "1234 5678"

# Function to replace text in runs while preserving formatting
def replace_text_in_paragraph(paragraph, placeholder, replacement):
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        # Replace the text in the full text
        full_text = full_text.replace(placeholder, replacement)
        
        # Clear existing runs and split the replacement text back into new runs
        for run in paragraph.runs:
            run.text = ''  # Clear the text in each run
        paragraph.runs[0].text = full_text  # Set the text in the first run

def insert_table_at_placeholder(doc, placeholder, case_details):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Clear the paragraph's text and insert the table
            paragraph.clear()  # Clear the placeholder text

            # Add a table at this location
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'  # Use a style of your choice

            # Define the header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Sagsnummer"
            header_cells[1].text = "Sagstitel"
            header_cells[2].text = "Sagsdato"
            header_cells[3].text = "Aktindsigt"

            # Add a row for each case
            for case_detail in case_details:
                row_cells = table.add_row().cells
                row_cells[0].text = case_detail[0]
                row_cells[1].text = case_detail[1]
                row_cells[2].text = case_detail[2]
                row_cells[3].text = case_detail[3]

            # Insert the table after clearing the placeholder
            paragraph._element.addnext(table._element)
            break


insert_table_at_placeholder(doc, "[Sagstabel]", case_details)

# Replace placeholders in paragraphs
for paragraph in doc.paragraphs:
    replace_text_in_paragraph(paragraph, '[Afdeling]', afdeling)
    replace_text_in_paragraph(paragraph, '[Ansøgernavn]', ansoegernavn)
    replace_text_in_paragraph(paragraph, '[Ansøgermail]', ansoegermail)
    replace_text_in_paragraph(paragraph, '[Dato]', dato)
    replace_text_in_paragraph(paragraph, '[Deskprotitel]', deskprotitel)
    replace_text_in_paragraph(paragraph, '[Besvarelse]', besvarelse)
    replace_text_in_paragraph(paragraph, '[Afdelingsmail]', afdelingsmail)
    replace_text_in_paragraph(paragraph, '[Afdelingstelefon]', afdelingstelefon)

# Replace placeholders in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, '[Afdeling]', afdeling)
                replace_text_in_paragraph(paragraph, '[Ansøgernavn]', ansoegernavn)
                replace_text_in_paragraph(paragraph, '[Ansøgermail]', ansoegermail)
                replace_text_in_paragraph(paragraph, '[Dato]', dato)
                replace_text_in_paragraph(paragraph, '[Deskprotitel]', deskprotitel)
                replace_text_in_paragraph(paragraph, '[Besvarelse]', besvarelse)
                replace_text_in_paragraph(paragraph, '[Afdelingsmail]', afdelingsmail)
                replace_text_in_paragraph(paragraph, '[Afdelingstelefon]', afdelingstelefon)

# Save the modified document
doc.save('ModifiedDocument.docx')
